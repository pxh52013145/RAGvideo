# -*- coding: utf-8 -*-
from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt


def main() -> None:
    root = Path(__file__).resolve().parents[1]
    docx_path = root / "project.docx"
    resource_dir = root / "resource"

    img_ui = resource_dir / "PixPin_2026-01-04_07-22-54.png"
    img_kb_flow = resource_dir / "知识库工作流.png"
    img_chat_flow = resource_dir / "dify工作流.png"

    for p in (docx_path, img_ui, img_kb_flow, img_chat_flow):
        if not p.exists():
            raise SystemExit(f"missing: {p}")

    doc = Document(str(docx_path))
    if not doc.tables:
        raise SystemExit("No table found in project.docx")

    table = doc.tables[0]
    if len(table.rows) < 5:
        raise SystemExit("Unexpected table shape: expected >= 5 rows")

    cell = table.cell(4, 0)

    # Clear existing placeholder content.
    cell.text = ""

    def set_run_font(run, size_pt: int | None = None, bold: bool | None = None) -> None:
        if bold is not None:
            run.bold = bold
        if size_pt is not None:
            run.font.size = Pt(size_pt)

    def add_heading(text: str) -> None:
        p = cell.add_paragraph()
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after = Pt(2)
        run = p.add_run(text)
        set_run_font(run, size_pt=14, bold=True)

    def add_subheading(text: str) -> None:
        p = cell.add_paragraph()
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after = Pt(1)
        run = p.add_run(text)
        set_run_font(run, size_pt=12, bold=True)

    def add_para(text: str) -> None:
        p = cell.add_paragraph(text)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)

    def add_bullets(lines: list[str], prefix: str = "• ") -> None:
        for line in lines:
            add_para(f"{prefix}{line}")

    def add_image(path: Path, caption: str, width_in: float = 5.8) -> None:
        p = cell.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(str(path), width=Inches(width_in))

        cap = cell.add_paragraph(caption)
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if cap.runs:
            set_run_font(cap.runs[0], size_pt=10, bold=False)

    # ---------------------------
    # Report Content
    # ---------------------------
    add_heading("项目背景")
    add_bullets(
        [
            "教学/培训视频信息密度高，但检索成本高：想找某个知识点往往需要反复拖动进度条。",
            "本项目将“视频 → 转录字幕 → 结构化笔记 → 知识库向量化 → 智能体问答”串成闭环，让用户用自然语言直接定位到视频片段并获得引用与时间戳。",
            "创新点：视频内容不仅生成 Markdown 笔记（可选），还会与原始字幕一起入库，支持跨视频的知识检索与对话式问答。",
            "竞争力：复用 Dify 的知识库、分片、Embedding、混合检索与工作流编排能力，降低从零实现 RAG 平台的成本，迭代更快。",
            "可扩展：未来可将时间戳与播放器联动，实现“点击引用跳转到对应时间点”。",
        ]
    )

    add_heading("相关工作")
    add_bullets(
        [
            "BiliNote：开源视频笔记项目，具备视频处理、字幕转录、截图与笔记生成等能力。本项目在其基础上引入“知识库入库 + RAG 问答”。",
            "Dify：开源 LLM 应用平台，提供数据集（知识库）、文档分片与向量索引、工作流编排、对话 API 与可视化运维能力。",
            "Whisper / faster-whisper：语音转文字（字幕/时间戳）核心组件，支持 CPU 模式运行。",
            "Ollama/云模型：Embedding 可本地部署（如 mxbai-embed-large、bge-m3 等），LLM 可对接官方 API（如 DeepSeek 等）或企业中转。",
            "对标：B站 AI 总结、Notion AI 等。优势在于“可私有化部署 + 视频时间戳引用 + 知识库可持续累积”。",
        ]
    )

    add_heading("设计方案")
    add_subheading("总体架构")
    add_bullets(
        [
            "客户端：Tauri + Vite 前端，提供入库、RAG 对话、引用展示等 UI。",
            "本地后端：FastAPI 负责视频处理（下载/读取、ffmpeg 抽取音频、Whisper 转录、可选生成 Markdown 笔记）、以及与 Dify API 的集成。",
            "Dify（自建）：提供知识库分片/向量化/检索与工作流（RAG Chat）执行。",
            "模型提供方：Embedding/LLM 可选择本地（Ollama）或云端 API（官方/中转）。",
        ]
    )

    add_subheading("入库工作流（Indexing）")
    add_bullets(
        [
            "输入：在线视频链接（如 B 站/抖音等）或本地视频文件。",
            "后端提取音频并转录得到带时间戳字幕；可选调用大模型生成结构化 Markdown 笔记。",
            "后端将“字幕 +（可选）笔记 + 元数据（视频标题/来源/时间戳段落）”组合成纯文本，调用 Dify Dataset API 写入知识库。",
            "Dify 根据数据集配置自动进行分片（chunk）、Embedding 与索引构建，形成可检索的向量库。",
        ]
    )

    add_subheading("对话工作流（RAG Chat）")
    add_bullets(
        [
            "用户问题从客户端提交到后端，后端转发到 Dify Chat/Workflow API。",
            "Dify 工作流内部执行知识库检索（向量/关键词混合检索），并把召回片段作为上下文供 LLM 生成回答。",
            "输出：回答文本 + 引用片段（包含对应视频/段落/时间戳），便于用户追溯原文。",
        ]
    )

    add_subheading("关键数据与可扩展点")
    add_bullets(
        [
            "入库文本保留时间戳字段（start/end），以“引用 = 视频 + 片段 + 时间范围”形式返回。",
            "后续可在客户端根据时间戳联动播放器（本地文件或在线播放页），实现点击跳转。",
            "支持保留 Markdown 笔记文件：便于离线阅读、二次编辑与导出。",
        ]
    )

    add_image(img_ui, "图1：客户端界面（入库 / RAG 对话）", width_in=5.8)
    add_image(img_kb_flow, "图2：知识库入库工作流（Dify Dataset Indexing）", width_in=5.8)
    add_image(img_chat_flow, "图3：对话工作流（Dify Workflow / Chat Messages）", width_in=5.8)

    add_heading("操作手册")
    add_subheading("1）Dify 侧（一次性配置）")
    add_bullets(
        [
            "使用 Docker 部署 Dify（按官方文档执行），在控制台添加 LLM 与 Embedding Provider。",
            "创建 Dataset（知识库），记录 Dataset ID。设置索引方式（high_quality/economy）与分片参数。",
            "创建 Chat App 或 Workflow App，并发布（Publish）。在“API 访问”中获取 App API Key；在“后端服务 API”获取 Service API Key。",
        ]
    )

    add_subheading("2）本项目侧（环境变量）")
    add_bullets(
        [
            "复制 `.env.example` 为 `.env`，填写：`DIFY_BASE_URL`、`DIFY_DATASET_ID`、`DIFY_SERVICE_API_KEY`、`DIFY_APP_API_KEY`。",
            "安装并配置 ffmpeg（或在 `.env` 中填写 `FFMPEG_BIN_PATH`）。",
        ]
    )

    add_subheading("3）启动与使用")
    add_bullets(
        [
            "启动后端：`python backend/main.py`（默认端口 8483）。",
            "启动前端：开发模式用 `pnpm dev`（浏览器访问），桌面端用 `pnpm tauri dev`。",
            "在“入库”页提交视频链接或本地文件，等待转录与（可选）笔记生成完成后自动写入 Dify 知识库。",
            "在“RAG 对话”页提问，系统返回答案与引用（含视频片段/时间戳）。",
            "可选加分：用 Cloudflare Tunnel / ngrok 将本机服务暴露为公网 HTTPS 地址用于验收演示。",
        ]
    )

    add_heading("项目总结")
    add_subheading("当前完成情况（截至中期）")
    add_bullets(
        [
            "完成“视频转录 + 可选笔记生成 + 自动入库到 Dify Dataset”的端到端流程。",
            "完成“Dify 工作流检索 + 对话 API + 引用返回”的 RAG 问答链路，并在客户端集成展示。",
            "支持本地部署与 Docker 化部署（前端/后端/nginx），便于在不同机器上快速启动。",
        ]
    )

    add_subheading("小组分工")
    add_bullets(
        [
            "成员1：后端（视频处理/转录/笔记生成）、Dify Dataset 写入与 Chat/Workflow API 集成、入库任务管理。",
            "成员2：前端（Tauri UI/交互/状态展示）、打包发布与跨平台兼容、文档与演示材料整理。",
            "共同：需求梳理、模型/参数调优、测试与验收演示。",
        ]
    )

    add_subheading("存在问题与后续规划")
    add_bullets(
        [
            "检索效果依赖分片与 Embedding/LLM 选择，下一阶段将优化：分片粒度、元数据结构与检索参数，并补充评测问题集。",
            "UI/交互：补齐对话历史、入库任务列表展示与错误提示，提升易用性。",
            "功能拓展：时间戳跳转播放（可选实现）、批量入库、多知识库管理与导出 Markdown 笔记。",
            "部署与加分项：完善一键启动/一键打包脚本，提供公网演示地址与 Docker 部署说明。",
        ]
    )

    doc.save(str(docx_path))
    print(f"updated {docx_path}")


if __name__ == "__main__":
    main()

